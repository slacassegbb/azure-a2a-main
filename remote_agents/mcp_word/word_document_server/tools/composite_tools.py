"""
Composite tools for Word Document Server.

These tools combine multiple operations into a single call to avoid
Azure AI Foundry's MCP chained tool call limitations.
"""
import os
import json
import shutil
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def build_document(
    filename: str,
    sections: List[Dict[str, Any]],
    title: Optional[str] = None,
    author: Optional[str] = None,
) -> str:
    """Create a complete Word document from a structured specification in a single call.

    This tool creates a new document and populates it with all content in one
    operation.  Use it instead of calling create_document + add_heading +
    add_paragraph etc. individually.

    Args:
        filename: Name for the document (with or without .docx extension).
        sections: Ordered list of content blocks. Each block is a dict with a
            ``type`` key and type-specific fields.  Supported types:

            - ``{"type": "heading", "text": "...", "level": 1, "font_name": null, "font_size": null, "bold": null, "italic": null, "border_bottom": false}``
            - ``{"type": "paragraph", "text": "...", "style": null, "font_name": null, "font_size": null, "bold": null, "italic": null, "color": null}``
            - ``{"type": "table", "rows": 3, "cols": 4, "data": [["a","b","c","d"], ...], "header_color": null, "text_color": null}``
            - ``{"type": "picture", "image_path": "...", "width": null, "source_type": "file"}``
            - ``{"type": "page_break"}``

        title: Optional document metadata title.
        author: Optional document metadata author.

    Returns:
        JSON string with status, filename, section count, and any per-section
        errors that occurred (the document is still saved even if individual
        sections fail).
    """
    filename = ensure_docx_extension(filename)

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return json.dumps({"error": f"Cannot create document: {error_message}"})

    try:
        doc = Document()

        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        ensure_heading_style(doc)
        ensure_table_style(doc)

        results: List[Dict[str, Any]] = []

        for idx, section in enumerate(sections):
            section_type = section.get("type", "").lower()
            try:
                if section_type == "heading":
                    _add_heading(doc, section)
                    results.append({"index": idx, "type": "heading", "ok": True})

                elif section_type == "paragraph":
                    _add_paragraph(doc, section)
                    results.append({"index": idx, "type": "paragraph", "ok": True})

                elif section_type == "table":
                    _add_table(doc, section)
                    results.append({"index": idx, "type": "table", "ok": True})

                elif section_type == "picture":
                    result = await _add_picture(doc, section)
                    results.append({"index": idx, "type": "picture", "ok": result is None, "error": result})

                elif section_type == "page_break":
                    doc.add_page_break()
                    results.append({"index": idx, "type": "page_break", "ok": True})

                else:
                    results.append({"index": idx, "type": section_type, "ok": False, "error": f"Unknown section type: {section_type}"})

            except Exception as e:
                results.append({"index": idx, "type": section_type, "ok": False, "error": str(e)})

        doc.save(filename)

        errors = [r for r in results if not r.get("ok")]

        # Prepare download
        download_dir = "/tmp/docx_downloads"
        os.makedirs(download_dir, exist_ok=True)
        safe_name = os.path.basename(filename).replace("/", "_").replace("\\", "_")
        if not safe_name.endswith(".docx"):
            safe_name += ".docx"
        dest_path = os.path.join(download_dir, safe_name)
        shutil.copy2(filename, dest_path)
        file_size = os.path.getsize(dest_path)

        return json.dumps({
            "message": "Document created and ready for download",
            "filename": safe_name,
            "sections_processed": len(results),
            "errors": errors if errors else None,
            "size_bytes": file_size,
            "download_url": f"/download/{safe_name}",
        })

    except Exception as e:
        return json.dumps({"error": f"Failed to build document: {str(e)}"})


# ── Internal helpers (not MCP tools) ──────────────────────────────────


def _add_heading(doc: Document, spec: Dict[str, Any]) -> None:
    text = spec.get("text", "")
    level = int(spec.get("level", 1))
    level = max(1, min(level, 9))

    try:
        heading = doc.add_heading(text, level=level)
    except Exception:
        heading = doc.add_paragraph(text)
        heading.style = doc.styles["Normal"]
        if heading.runs:
            run = heading.runs[0]
            run.bold = True
            run.font.size = Pt({1: 16, 2: 14}.get(level, 12))

    _apply_run_formatting(heading, spec)

    if spec.get("border_bottom"):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = heading._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)


def _add_paragraph(doc: Document, spec: Dict[str, Any]) -> None:
    text = spec.get("text", "")
    paragraph = doc.add_paragraph(text)

    style = spec.get("style")
    if style:
        try:
            paragraph.style = style
        except KeyError:
            paragraph.style = doc.styles["Normal"]

    _apply_run_formatting(paragraph, spec)


def _add_table(doc: Document, spec: Dict[str, Any]) -> None:
    rows = int(spec.get("rows", 1))
    cols = int(spec.get("cols", 1))
    data = spec.get("data")

    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    if data:
        for i, row_data in enumerate(data):
            if i >= rows:
                break
            for j, cell_text in enumerate(row_data):
                if j >= cols:
                    break
                table.cell(i, j).text = str(cell_text)

    # Optional header highlighting
    header_color = spec.get("header_color")
    text_color = spec.get("text_color", "FFFFFF")
    if header_color and rows > 0:
        from word_document_server.core.tables import highlight_header_row
        highlight_header_row(table, header_color, text_color)


async def _add_picture(doc: Document, spec: Dict[str, Any]) -> Optional[str]:
    """Returns None on success, error string on failure."""
    import tempfile
    import urllib.request

    image_path = spec.get("image_path", "")
    width = spec.get("width")
    source_type = spec.get("source_type", "file")

    temp_path = None
    try:
        if source_type == "url" or image_path.startswith(("http://", "https://")):
            url_path = image_path.split("?")[0]
            ext = os.path.splitext(url_path)[1] or ".png"
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                temp_path = tmp.name
            urllib.request.urlretrieve(image_path, temp_path)
            abs_path = temp_path
        else:
            abs_path = os.path.abspath(image_path)
            if not os.path.exists(abs_path):
                return f"Image not found: {abs_path}"

        if width:
            doc.add_picture(abs_path, width=Inches(float(width)))
        else:
            doc.add_picture(abs_path)
        return None
    except Exception as e:
        return str(e)
    finally:
        if temp_path and os.path.exists(temp_path):
            os.unlink(temp_path)


def _apply_run_formatting(paragraph, spec: Dict[str, Any]) -> None:
    font_name = spec.get("font_name")
    font_size = spec.get("font_size")
    bold = spec.get("bold")
    italic = spec.get("italic")
    color = spec.get("color")

    if not any(v is not None for v in [font_name, font_size, bold, italic, color]):
        return

    for run in paragraph.runs:
        if font_name:
            run.font.name = font_name
        if font_size is not None:
            run.font.size = Pt(int(font_size))
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color:
            color_hex = str(color).lstrip("#")
            try:
                run.font.color.rgb = RGBColor.from_string(color_hex)
            except Exception:
                pass
